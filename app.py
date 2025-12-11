"""
Submittal Extraction API
FastAPI application for extracting submittals from construction spec books
"""
import shutil

from fastapi import Request
from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
import logging
import time, datetime
from pathlib import Path
from io import BytesIO
import zipfile
from fastapi.responses import JSONResponse
import base64, tempfile

import os
from dotenv import load_dotenv
import json
from fastapi import HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from openai import OpenAI
import fitz  # PyMuPDF
from datetime import datetime

from extractor import SubmittalExtractor

load_dotenv()
# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Initialize FastAPI app
app = FastAPI(
    title="Submittal Extraction API",
    description="Extract submittal requirements from construction specification books",
    version="1.0.0",
    docs_url="/docs",
    redoc_url="/redoc"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure appropriately for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

OUTPUT_DIR = "./output"
TEMPLATES_DIR = "./templates"
os.makedirs(OUTPUT_DIR, exist_ok=True)

@app.get("/")
async def root():
    """
    Root endpoint - API information
    """
    return {
        "message": "Submittal Extraction API",
        "version": "1.0.0",
        "endpoints": {
            "health": "/health",
            "extract_excel": "/extract-submittals (returns ZIP with Excel files)",
            "extract_json": "/extract-submittals-json (returns JSON for n8n)",
            "docs": "/docs"
        }
    }


@app.get("/health")
async def health_check():
    """
    Health check endpoint
    Used by Azure and monitoring services
    """
    return {
        "status": "healthy",
        "service": "submittal-extraction-api",
        "version": "1.0.0",
        "timestamp": time.time()
    }


@app.post("/extract-submittals-sections")
async def extract_submittals_sections(request: Request):
    """
    Extract submittal sections from PDF and return submittal_sections.xlsx
    Accepts base64 encoded PDF file
    Returns: Excel file with submittal sections
    """
    start_time = time.time()
    temp_file_path = None

    try:
        # Parse JSON manually
        body = await request.json()

        filename = body.get("filename")
        file_content = body.get("file_content")

        logger.info(f"[SECTIONS] Received filename: {filename}")
        logger.info(f"[SECTIONS] Received file_content length: {len(file_content) if file_content else 0}")

        if not file_content:
            raise HTTPException(400, "file_content is required")

        if not filename:
            filename = "document.pdf"

        if not filename.lower().endswith('.pdf'):
            raise HTTPException(400, "Invalid file type")

        # Decode base64
        import base64
        file_bytes = base64.b64decode(file_content)
        logger.info(f"[SECTIONS] Decoded file size: {len(file_bytes)} bytes")

        # Validate PDF header
        if file_bytes[:4] != b'%PDF':
            raise HTTPException(400, f"Invalid PDF - first bytes: {file_bytes[:4]}")

        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf', mode='wb') as temp_file:
            temp_file_path = temp_file.name
            temp_file.write(file_bytes)
            temp_file.flush()
            logger.info(f"[SECTIONS] File saved to: {temp_file_path}")

        # Look for template
        template_path = Path(__file__).parent / "templates" / "SubmittalLog.xlsx"
        if not template_path.exists():
            template_path = Path("templates/SubmittalLog.xlsx")
            if not template_path.exists():
                template_path = None

        # Extract
        logger.info(f"[SECTIONS] Starting extraction...")
        extractor = SubmittalExtractor(temp_file_path)
        result = extractor.extract(template_path=str(template_path) if template_path else None)

        extraction_time = time.time() - start_time
        sections_wb = result["sections"]

        logger.info(f"[SECTIONS] Extraction completed: {len(sections_wb.sheetnames)} sheets, {extraction_time:.2f}s")

        # Save sections Excel to BytesIO
        sections_excel = BytesIO()
        sections_wb.save(sections_excel)
        sections_excel.seek(0)

        # Create filename
        pdf_name = Path(filename).stem
        excel_filename = f"{pdf_name}_submittal_sections.xlsx"

        logger.info(f"[SECTIONS] Returning Excel: {excel_filename}")

        return StreamingResponse(
            sections_excel,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={excel_filename}"}
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[SECTIONS] Error: {str(e)}", exc_info=True)
        raise HTTPException(500, f"Extraction failed: {str(e)}")
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
                logger.info(f"[SECTIONS] Cleaned up temp file: {temp_file_path}")
            except Exception as e:
                logger.warning(f"[SECTIONS] Cleanup failed: {str(e)}")


@app.post("/extract-submittals-log")
async def extract_submittals_log(request: Request):
    """
    Extract submittal log from PDF and return submittals_log.xlsx
    Accepts base64 encoded PDF file
    Returns: Excel file with submittal log
    """
    start_time = time.time()
    temp_file_path = None

    try:
        # Parse JSON manually
        body = await request.json()

        filename = body.get("filename")
        file_content = body.get("file_content")

        logger.info(f"[LOG] Received filename: {filename}")
        logger.info(f"[LOG] Received file_content length: {len(file_content) if file_content else 0}")

        if not file_content:
            raise HTTPException(400, "file_content is required")

        if not filename:
            filename = "document.pdf"

        if not filename.lower().endswith('.pdf'):
            raise HTTPException(400, "Invalid file type")

        # Decode base64
        import base64
        file_bytes = base64.b64decode(file_content)
        logger.info(f"[LOG] Decoded file size: {len(file_bytes)} bytes")

        # Validate PDF header
        if file_bytes[:4] != b'%PDF':
            raise HTTPException(400, f"Invalid PDF - first bytes: {file_bytes[:4]}")

        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf', mode='wb') as temp_file:
            temp_file_path = temp_file.name
            temp_file.write(file_bytes)
            temp_file.flush()
            logger.info(f"[LOG] File saved to: {temp_file_path}")

        # Look for template
        template_path = Path(__file__).parent / "templates" / "SubmittalLog.xlsx"
        if not template_path.exists():
            template_path = Path("templates/SubmittalLog.xlsx")
            if not template_path.exists():
                template_path = None

        # Extract
        logger.info(f"[LOG] Starting extraction...")
        extractor = SubmittalExtractor(temp_file_path)
        result = extractor.extract(template_path=str(template_path) if template_path else None)

        extraction_time = time.time() - start_time
        log_wb = result["log"]

        logger.info(f"[LOG] Extraction completed in {extraction_time:.2f}s")

        # Save log Excel to BytesIO
        log_excel = BytesIO()
        log_wb.save(log_excel)
        log_excel.seek(0)

        # Create filename
        pdf_name = Path(filename).stem
        excel_filename = f"{pdf_name}_submittals_log.xlsx"

        logger.info(f"[LOG] Returning Excel: {excel_filename}")

        return StreamingResponse(
            log_excel,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={excel_filename}"}
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[LOG] Error: {str(e)}", exc_info=True)
        raise HTTPException(500, f"Extraction failed: {str(e)}")
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
                logger.info(f"[LOG] Cleaned up temp file: {temp_file_path}")
            except Exception as e:
                logger.warning(f"[LOG] Cleanup failed: {str(e)}")


class ProjectInfoRequest(BaseModel):
    filename: str
    file_content: str


class SubmittalStructureRequest(BaseModel):
    filename: str
    file_content: str
    project_info: dict


# ===== API ENDPOINT 3: EXTRACT PROJECT INFO =====

@app.post("/extract-project-info")
async def extract_project_info(request: ProjectInfoRequest):
    """Extract project information from first 10 pages of PDF using OpenAI"""
    temp_file_path = None

    try:
        logger.info(f"Extracting project info from: {request.filename}")

        # Get OpenAI API key
        openai_api_key = os.getenv("OPENAI_API_KEY")
        if not openai_api_key:
            raise HTTPException(500, "OPENAI_API_KEY not configured")

        # Decode base64
        file_bytes = base64.b64decode(request.file_content)
        logger.info(f"Decoded file size: {len(file_bytes)} bytes")

        if file_bytes[:4] != b'%PDF':
            raise HTTPException(400, "Invalid PDF file")

        # Save to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf', mode='wb') as temp_file:
            temp_file_path = temp_file.name
            temp_file.write(file_bytes)
            temp_file.flush()

        # Extract text from first 10 pages
        doc = fitz.open(temp_file_path)
        total_pages = len(doc)
        pages_to_extract = min(10, total_pages)

        pdf_text = ""
        for page_num in range(pages_to_extract):
            try:
                page = doc[page_num]
                text = page.get_text("text", sort=True)
                pdf_text += f"\n--- Page {page_num + 1} ---\n{text}"
            except Exception as e:
                logger.warning(f"Error extracting page {page_num + 1}: {e}")

        doc.close()
        pdf_text = pdf_text[:12000]

        # Prepare OpenAI prompt
        prompt = f"""Extract the following project info (even if wording varies or formatting is inconsistent):

- Project Name (also appears as "Project Title", "Project", or on title page)
- Project Number (also called "Project No.", "Contract No.", "CCUA Project Number", "Project ID")
- PSCC Job Number (also appears as "PSCC Project Number", "Job No.", "Job Number")
- Engineer Name and Address (may appear as a firm name, contact block, or footer stamp)
- Contractor Name and Address (may appear on signature page or contract information sheet)
- Owner Name and Address (may appear as "Owner", "Client", "Utility Authority", etc.)
- Prepared By (may be an engineer, consultant, or preparation firm listed on cover/title page)

IMPORTANT:
- If a value appears multiple times, use the most complete version.
- If a field is partially available (e.g., name but not address), populate what is available and leave the missing part empty.
- If a field is truly not present, return an empty string "" for that key.
- Never omit any key.

Return ONLY strict JSON with keys: project_name, ccua_project_number, pscc_job_number, engineer_name, engineer_address, contractor_name, contractor_address, owner_name, owner_address, prepared_by

If a field is not found, use an empty string "".

PDF Text (first 10 pages):
\"\"\"
{pdf_text}
\"\"\"
"""

        # Call OpenAI
        client = OpenAI(api_key=openai_api_key)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system",
                 "content": "You are a helpful assistant that extracts project information from construction PDFs. Always return valid JSON only."},
                {"role": "user", "content": prompt}
            ],
            temperature=0,
            max_tokens=500
        )

        # Parse response
        openai_response = response.choices[0].message.content.strip()
        if openai_response.startswith("```json"):
            openai_response = openai_response.replace("```json", "").replace("```", "").strip()
        elif openai_response.startswith("```"):
            openai_response = openai_response.replace("```", "").strip()

        try:
            project_info = json.loads(openai_response)
        except json.JSONDecodeError as e:
            project_info = {"error": "Failed to parse response", "raw_response": openai_response}

        return {
            "success": True,
            "filename": request.filename,
            "pages_analyzed": pages_to_extract,
            "project_info": project_info
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error: {str(e)}", exc_info=True)
        raise HTTPException(500, f"Extraction failed: {str(e)}")
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
            except:
                pass


def replace_placeholder_in_paragraph(para, replacements):
    """
    Replace placeholders in a paragraph that may be split across runs

    Args:
        para: docx paragraph object
        replacements: dict of {placeholder: value}
    """
    # Get full paragraph text
    full_text = para.text

    # Check if any placeholder exists in full text
    needs_replacement = False
    for placeholder in replacements.keys():
        if placeholder in full_text:
            needs_replacement = True
            break

    if not needs_replacement:
        return False

    # Replace all placeholders in full text
    new_text = full_text
    for placeholder, value in replacements.items():
        if placeholder in new_text:
            new_text = new_text.replace(placeholder, str(value or ""))

    # Clear all runs and create a single new run with replaced text
    # Preserve formatting from first run
    if para.runs:
        first_run = para.runs[0]
        # Clear all runs
        for run in para.runs:
            run.text = ""
        # Set new text in first run
        first_run.text = new_text
    else:
        # No runs, add new one
        para.add_run(new_text)

    return True


def replace_in_table(table, replacements):
    """
    Recursively replace placeholders in a table, including nested tables

    Args:
        table: docx table object
        replacements: dict of {placeholder: value}
    """
    for row in table.rows:
        for cell in row.cells:
            # Replace in paragraphs
            for para in cell.paragraphs:
                replace_placeholder_in_paragraph(para, replacements)

            # Recursively process nested tables
            for nested_table in cell.tables:
                replace_in_table(nested_table, replacements)


# ===== API ENDPOINT 4: SUBMITTAL FOLDERS =====
@app.post("/create-submittal-structure")
async def create_submittal_structure_working(request: SubmittalStructureRequest):
    """
    Create folder structure with all placeholders correctly replaced
    Handles split placeholders and adds submittal name to Excel
    """
    temp_file_path = None
    temp_dir = None

    try:
        logger.info(f"Creating submittal structure: {request.filename}")

        # Decode and validate base64
        import base64
        file_bytes = base64.b64decode(request.file_content)
        if file_bytes[:4] != b'%PDF':
            raise HTTPException(400, "Invalid PDF file")

        # Save PDF
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf', mode='wb') as temp_file:
            temp_file_path = temp_file.name
            temp_file.write(file_bytes)
            temp_file.flush()

        # Use extractor to find sections with submittals
        from extractor import SubmittalExtractor
        import fitz

        extractor = SubmittalExtractor(temp_file_path)
        extractor.doc = fitz.open(temp_file_path)
        total_pages = len(extractor.doc)

        logger.info(f"PDF has {total_pages} pages")

        # Extract FULL PDF text
        logger.info("Extracting full PDF text...")
        extractor.full_text = ""
        for page_num in range(total_pages):
            try:
                page = extractor.doc[page_num]
                text = page.get_text("text", sort=True)
                extractor.full_text += text + "\n"
            except:
                try:
                    text = page.get_text("text", sort=False)
                    extractor.full_text += text + "\n"
                except:
                    pass

        # Parse TOC
        extractor._compile_patterns()
        extractor.toc = extractor._extract_toc_from_first_100_pages()

        if len(extractor.toc) == 0:
            extractor.toc = extractor._scan_pdf_for_sections()

        logger.info(f"Found {len(extractor.toc)} sections in TOC")

        # Process sections to find those with submittals
        extractor.sections_with_submittals = []
        for section_num_display, section_num_search, section_name in extractor.toc:
            extractor._process_section(section_num_display, section_num_search, section_name)

        extractor.doc.close()

        logger.info(f"Found {len(extractor.sections_with_submittals)} sections with submittals")

        if len(extractor.sections_with_submittals) == 0:
            raise HTTPException(400, "No sections with submittal subsections found")

        # Create temp directory
        temp_dir = tempfile.mkdtemp()
        base_path = Path(temp_dir)

        # Load templates
        word_template = Path(__file__).parent / "templates" / "Coverpage_New.docx"
        excel_template = Path(__file__).parent / "templates" / "Transmittal.xlsx"

        if not word_template.exists():
            raise HTTPException(500, "Coverpage_New.docx not found")
        if not excel_template.exists():
            raise HTTPException(500, "Transmittal.xlsx not found")

        # Create folders for each section with submittals
        created_count = 0

        for section_num, section_name in extractor.sections_with_submittals:
            try:
                # Clean folder name
                folder_name = f"{section_num} {section_name}"
                folder_name = "".join(c for c in folder_name if c.isalnum() or c in (' ', '-', '_', '.'))
                if len(folder_name) > 200:
                    folder_name = folder_name[:200]

                section_path = base_path / folder_name
                section_path.mkdir(parents=True, exist_ok=True)

                # Create subfolder structure
                vendor = section_path / "1. From Vendor"
                to_eng = section_path / "2. To Engineer"
                frm_eng = section_path / "3. From Engineer"
                final = section_path / "4. Final Approved"

                for folder in [vendor, to_eng, frm_eng, final]:
                    folder.mkdir(parents=True, exist_ok=True)

                # Create Rev folders
                for rev in ["Rev 0", "Rev 1", "Rev 2"]:
                    (to_eng / rev).mkdir(parents=True, exist_ok=True)

                # Fill Word document
                from docx import Document
                doc = Document(word_template)

                # Prepare replacements
                today_date = datetime.now().strftime("%m/%d/%Y")

                ccua_no = (request.project_info.get("ccua_project_number") or
                           request.project_info.get("project_number") or "")

                replacements = {
                    "{{PROJECT_NAME}}": request.project_info.get("project_name", ""),
                    "{{CCUA_PROJECT_NO}}": ccua_no,
                    "{{PSCC_JOB_NO}}": request.project_info.get("pscc_job_number", ""),
                    "{{ENGINEER_NAME}}": request.project_info.get("engineer_name", ""),
                    "{{ENGINEER_ADDRESS}}": request.project_info.get("engineer_address", ""),
                    "{{CONTRACTOR_NAME}}": request.project_info.get("contractor_name", ""),
                    "{{CONTRACTOR_ADDRESS}}": request.project_info.get("contractor_address", ""),
                    "{{OWNER_NAME}}": request.project_info.get("owner_name", ""),
                    "{{OWNER_ADDRESS}}": request.project_info.get("owner_address", ""),
                    "{{PREPARED_BY}}": request.project_info.get("prepared_by", ""),
                    "{{SECTION_NUMBER}}": section_num,
                    "{{SUBMITTAL_TITLE}}": section_name,
                    "{{DATE}}": today_date,
                }

                logger.debug(f"Processing section {section_num}: {section_name}")

                # Replace in all paragraphs
                for para in doc.paragraphs:
                    replace_placeholder_in_paragraph(para, replacements)

                # Replace in all tables (including nested tables)
                for table in doc.tables:
                    replace_in_table(table, replacements)

                # Replace in headers
                for section in doc.sections:
                    for para in section.header.paragraphs:
                        replace_placeholder_in_paragraph(para, replacements)

                # Replace in footers
                for section in doc.sections:
                    for para in section.footer.paragraphs:
                        replace_placeholder_in_paragraph(para, replacements)

                doc.save(to_eng / "Submittal_CoverPage.docx")
                logger.info(f"Saved Word doc for {section_num}")

                # Fill Excel document
                import openpyxl
                wb = openpyxl.load_workbook(excel_template)
                ws = wb.active

                # ADD SUBMITTAL NAME IN FIRST ROW
                # Insert a new row at the top
                ws.insert_rows(1)
                # Add submittal name
                ws['A1'] = f"Submittal for Section {section_num} - {section_name}"
                # Make it bold
                ws['A1'].font = openpyxl.styles.Font(bold=True, size=12)

                excel_replacements = {
                    "{{ENGINEER_NAME}}": request.project_info.get("engineer_name", ""),
                    "{{ENGINEER_ADDRESS}}": request.project_info.get("engineer_address", ""),
                    "{{OWNER_NAME}}": request.project_info.get("owner_name", ""),
                    "{{OWNER_ADDRESS}}": request.project_info.get("owner_address", ""),
                    "{{DATE}}": today_date,
                    "{{PSCC_JOB_NO}}": request.project_info.get("pscc_job_number", ""),
                    "{{PROJECT_NAME}}": request.project_info.get("project_name", ""),
                    "{{CCUA_PROJECT_NO}}": ccua_no,
                    "{{SECTION_NUMBER}}": section_num,
                    "{{SUBMITTAL_TITLE}}": section_name
                }

                # Replace placeholders in all cells
                for row in ws.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            for placeholder, value in excel_replacements.items():
                                if placeholder in cell.value:
                                    cell.value = cell.value.replace(placeholder, str(value or ""))

                wb.save(to_eng / "Transmittal.xlsx")
                logger.info(f"Saved Excel for {section_num}")

                created_count += 1

            except Exception as e:
                logger.error(f"Failed for {section_num}: {e}", exc_info=True)
                continue

        logger.info(f"Created {created_count} section folders")

        if created_count == 0:
            raise HTTPException(500, "Failed to create any folders")

        # Create ZIP file
        logger.info("Creating ZIP file...")
        zip_buffer = BytesIO()

        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    file_path = Path(root) / file
                    arcname = file_path.relative_to(temp_dir)
                    zip_file.write(file_path, arcname)

                for dir_name in dirs:
                    dir_path = Path(root) / dir_name
                    arcname = str(dir_path.relative_to(temp_dir)) + '/'
                    zip_info = zipfile.ZipInfo(arcname)
                    zip_info.external_attr = 0o040755 << 16
                    zip_file.writestr(zip_info, '')

        zip_buffer.seek(0)

        pdf_name = Path(request.filename).stem
        zip_filename = f"{pdf_name}_submittal_structure.zip"

        logger.info(f"ZIP ready: {zip_filename} ({created_count} sections)")

        return StreamingResponse(
            zip_buffer,
            media_type="application/zip",
            headers={"Content-Disposition": f"attachment; filename={zip_filename}"}
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error: {str(e)}", exc_info=True)
        raise HTTPException(500, f"Failed: {str(e)}")
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
            except:
                pass
        if temp_dir and os.path.exists(temp_dir):
            try:
                import shutil
                shutil.rmtree(temp_dir)
            except:
                pass


@app.get("/test")
async def test_endpoint():
    """
    Simple test endpoint to verify API is responding
    """
    return {
        "status": "API is working",
        "message": "Test successful",
        "timestamp": time.time()
    }


# Error handlers
@app.exception_handler(404)
async def not_found_handler(request, exc):
    return JSONResponse(
        status_code=404,
        content={"error": "Not found", "detail": "The requested resource was not found"}
    )


@app.exception_handler(500)
async def internal_error_handler(request, exc):
    logger.error(f"Internal server error: {str(exc)}", exc_info=True)
    return JSONResponse(
        status_code=500,
        content={"error": "Internal server error", "detail": "An unexpected error occurred"}
    )


# Startup event
@app.on_event("startup")
async def startup_event():
    logger.info("Submittal Extraction API starting up...")
    logger.info("API is ready to accept requests")


# Shutdown event
@app.on_event("shutdown")
async def shutdown_event():
    logger.info("Submittal Extraction API shutting down...")



if __name__ == "__main__":
    import uvicorn

    # Get port from environment variable (Azure uses PORT)
    port = int(os.getenv("PORT", 8000))

    uvicorn.run(
        "app:app",
        host="0.0.0.0",
        port=port,
        reload=True,  # Set to False in production
        log_level="info"
    )